import streamlit as st
import requests, smtplib, os, datetime, math, io, base64
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table,
    TableStyle, HRFlowable, Image as RLImage, PageBreak
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.graphics.shapes import Drawing, Line, Rect, String, Circle

# Word doc
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

from scoring_keys import (
    L_SCALE, F_SCALE, FB_SCALE, FP_SCALE, K_SCALE, S_SCALE,
    K_CORRECTIONS,
    HS_SCALE, D_SCALE, HY_SCALE, PD_SCALE, MF_MALE, MF_FEMALE,
    PA_SCALE, PT_SCALE, SC_SCALE, MA_SCALE, SI_SCALE,
    ANX_SCALE, FRS_SCALE, OBS_SCALE, DEP_SCALE, HEA_SCALE, BIZ_SCALE,
    ANG_SCALE, CYN_SCALE, ASP_SCALE, TPA_SCALE, LSE_SCALE, SOD_SCALE,
    FAM_SCALE, WRK_SCALE, TRT_SCALE,
    AGGR_SCALE, PSYC_SCALE, DISC_SCALE, NEGE_SCALE, INTR_SCALE,
    A_SCALE, R_SCALE, MACR_SCALE, ES_SCALE, DO_SCALE, RE_SCALE,
    MT_SCALE, OH_SCALE, APS_SCALE, AAS_SCALE, MDS_SCALE, HO_SCALE, PK_SCALE,
    GM_SCALE, GF_SCALE,
    D1_SUB, D2_SUB, D3_SUB, D4_SUB, D5_SUB,
    HY1_SUB, HY2_SUB, HY3_SUB, HY4_SUB, HY5_SUB,
    PD1_SUB, PD2_SUB, PD3_SUB, PD4_SUB, PD5_SUB,
    PA1_SUB, PA2_SUB, PA3_SUB,
    SC1_SUB, SC2_SUB, SC3_SUB, SC4_SUB, SC5_SUB, SC6_SUB,
    MA1_SUB, MA2_SUB, MA3_SUB, MA4_SUB,
    SI1_SUB, SI2_SUB, SI3_SUB,
    NORMATIVE_DATA, CRITICAL_ITEMS_KB, CRITICAL_ITEMS_LW,
    VRIN_PAIRS, TRIN_TRUE_PAIRS, TRIN_FALSE_PAIRS,
)

from arabic_questions import MMPI2_QUESTIONS_AR

# ══════════════════════════════════════════════════════════════
#  CONFIGURATION
# ══════════════════════════════════════════════════════════════

GMAIL_ADDRESS   = "Wijdan.psyc@gmail.com"
GMAIL_PASSWORD  = "rias eeul lyuu stce"
THERAPIST_EMAIL = "Wijdan.psyc@gmail.com"
LOGO_FILE       = "logo.png"
ITEMS_PER_PAGE  = 50

THERAPIST_NAME  = "Yusuf Abdelatti"
THERAPIST_TITLE = "Psychotherapist"
CENTER_NAME     = "Wijdan Therapy Center"

# ══════════════════════════════════════════════════════════════
#  SCORING ENGINE (unchanged)
# ══════════════════════════════════════════════════════════════

def score_scale(responses, scale_key):
    return sum(1 for item, keyed in scale_key.items()
               if item in responses and responses[item] == keyed)

def raw_to_t(raw, scale_name, gender):
    if raw is None or scale_name not in NORMATIVE_DATA:
        return 50
    norm = NORMATIVE_DATA[scale_name]
    if len(norm) < 4:
        return 50
    m, sd, f_m, f_sd = norm
    mean = m if gender == "Male" else f_m
    sd_v = sd if gender == "Male" else f_sd
    if not sd_v:
        return 50
    return max(20, min(120, round(50 + 10 * (raw - mean) / sd_v)))

def compute_vrin(responses):
    score = 0
    for pair in VRIN_PAIRS:
        i1, i2, pattern = pair
        if i1 in responses and i2 in responses:
            r1, r2 = responses[i1], responses[i2]
            if pattern == "TF" and r1 is True  and r2 is False: score += 1
            elif pattern == "FT" and r1 is False and r2 is True:  score += 1
    return score

def compute_trin(responses):
    score = 9
    for i1, i2 in TRIN_TRUE_PAIRS:
        if i1 in responses and i2 in responses:
            if responses[i1] is True and responses[i2] is True:
                score += 1
    for i1, i2 in TRIN_FALSE_PAIRS:
        if i1 in responses and i2 in responses:
            if responses[i1] is False and responses[i2] is False:
                score -= 1
    return max(0, score)

def compute_all_scores(responses, gender):
    r = responses
    cannot_say = sum(1 for v in r.values() if v is None)
    l_raw  = score_scale(r, L_SCALE)
    f_raw  = score_scale(r, F_SCALE)
    fb_raw = score_scale(r, FB_SCALE)
    fp_raw = score_scale(r, FP_SCALE)
    k_raw  = score_scale(r, K_SCALE)
    s_raw  = score_scale(r, S_SCALE)
    vrin_raw = compute_vrin(r)
    trin_raw = compute_trin(r)

    hs_raw = score_scale(r, HS_SCALE)
    d_raw  = score_scale(r, D_SCALE)
    hy_raw = score_scale(r, HY_SCALE)
    pd_raw = score_scale(r, PD_SCALE)
    mf_key = MF_MALE if gender == "Male" else MF_FEMALE
    mf_raw = score_scale(r, mf_key)
    pa_raw = score_scale(r, PA_SCALE)
    pt_raw = score_scale(r, PT_SCALE)
    sc_raw = score_scale(r, SC_SCALE)
    ma_raw = score_scale(r, MA_SCALE)
    si_raw = score_scale(r, SI_SCALE)

    hs_k = hs_raw + round(K_CORRECTIONS["Hs"] * k_raw)
    pd_k = pd_raw + round(K_CORRECTIONS["Pd"] * k_raw)
    pt_k = pt_raw + round(K_CORRECTIONS["Pt"] * k_raw)
    sc_k = sc_raw + round(K_CORRECTIONS["Sc"] * k_raw)
    ma_k = ma_raw + round(K_CORRECTIONS["Ma"] * k_raw)

    def ct(raw, name): return raw_to_t(raw, name, gender)

    scores = {
        "cannot_say": cannot_say,
        "L_raw": l_raw, "F_raw": f_raw, "Fb_raw": fb_raw,
        "Fp_raw": fp_raw, "K_raw": k_raw, "S_raw": s_raw,
        "VRIN_raw": vrin_raw, "TRIN_raw": trin_raw,
        "L_T": ct(l_raw,"L"), "F_T": ct(f_raw,"F"),
        "Fb_T": ct(fb_raw,"Fb"), "Fp_T": ct(fp_raw,"Fp"),
        "K_T": ct(k_raw,"K"), "S_T": ct(s_raw,"S"),
        "VRIN_T": ct(vrin_raw,"VRIN"), "TRIN_T": ct(trin_raw,"TRIN"),
        "Hs_raw": hs_raw, "D_raw": d_raw, "Hy_raw": hy_raw,
        "Pd_raw": pd_raw, "Mf_raw": mf_raw, "Pa_raw": pa_raw,
        "Pt_raw": pt_raw, "Sc_raw": sc_raw, "Ma_raw": ma_raw, "Si_raw": si_raw,
        "Hs_k": hs_k, "Pd_k": pd_k, "Pt_k": pt_k, "Sc_k": sc_k, "Ma_k": ma_k,
        "Hs_T": ct(hs_k,"Hs"), "D_T": ct(d_raw,"D"),
        "Hy_T": ct(hy_raw,"Hy"), "Pd_T": ct(pd_k,"Pd"),
        "Mf_T": ct(mf_raw,"Mf_M" if gender=="Male" else "Mf_F"),
        "Pa_T": ct(pa_raw,"Pa"), "Pt_T": ct(pt_k,"Pt"),
        "Sc_T": ct(sc_k,"Sc"), "Ma_T": ct(ma_k,"Ma"), "Si_T": ct(si_raw,"Si"),
        "FK_index": f_raw - k_raw,
    }

    for name, key in [
        ("ANX",ANX_SCALE),("FRS",FRS_SCALE),("OBS",OBS_SCALE),
        ("DEP",DEP_SCALE),("HEA",HEA_SCALE),("BIZ",BIZ_SCALE),
        ("ANG",ANG_SCALE),("CYN",CYN_SCALE),("ASP",ASP_SCALE),
        ("TPA",TPA_SCALE),("LSE",LSE_SCALE),("SOD",SOD_SCALE),
        ("FAM",FAM_SCALE),("WRK",WRK_SCALE),("TRT",TRT_SCALE),
    ]:
        raw = score_scale(r, key)
        scores[f"{name}_raw"] = raw; scores[f"{name}_T"] = ct(raw, name)

    for name, key in [
        ("AGGR",AGGR_SCALE),("PSYC",PSYC_SCALE),("DISC",DISC_SCALE),
        ("NEGE",NEGE_SCALE),("INTR",INTR_SCALE),
    ]:
        raw = score_scale(r, key)
        scores[f"{name}_raw"] = raw; scores[f"{name}_T"] = ct(raw, name)

    for name, key in [
        ("A",A_SCALE),("R",R_SCALE),("MAC_R",MACR_SCALE),
        ("Es",ES_SCALE),("Do",DO_SCALE),("Re",RE_SCALE),
        ("Mt",MT_SCALE),("OH",OH_SCALE),("APS",APS_SCALE),
        ("AAS",AAS_SCALE),("MDS",MDS_SCALE),("Ho",HO_SCALE),
        ("PK",PK_SCALE),("GM",GM_SCALE),("GF",GF_SCALE),
    ]:
        raw = score_scale(r, key)
        scores[f"{name}_raw"] = raw; scores[f"{name}_T"] = ct(raw, name)

    for name, key in [
        ("D1",D1_SUB),("D2",D2_SUB),("D3",D3_SUB),("D4",D4_SUB),("D5",D5_SUB),
        ("Hy1",HY1_SUB),("Hy2",HY2_SUB),("Hy3",HY3_SUB),("Hy4",HY4_SUB),("Hy5",HY5_SUB),
        ("Pd1",PD1_SUB),("Pd2",PD2_SUB),("Pd3",PD3_SUB),("Pd4",PD4_SUB),("Pd5",PD5_SUB),
        ("Pa1",PA1_SUB),("Pa2",PA2_SUB),("Pa3",PA3_SUB),
        ("Sc1",SC1_SUB),("Sc2",SC2_SUB),("Sc3",SC3_SUB),
        ("Sc4",SC4_SUB),("Sc5",SC5_SUB),("Sc6",SC6_SUB),
        ("Ma1",MA1_SUB),("Ma2",MA2_SUB),("Ma3",MA3_SUB),("Ma4",MA4_SUB),
        ("Si1",SI1_SUB),("Si2",SI2_SUB),("Si3",SI3_SUB),
    ]:
        raw = score_scale(r, key)
        scores[f"{name}_raw"] = raw; scores[f"{name}_T"] = ct(raw, name)

    critical_kb = {}
    for cat, items in CRITICAL_ITEMS_KB.items():
        flagged = [i for i in items if i in r and r[i] is True]
        if flagged: critical_kb[cat] = flagged
    scores["critical_kb"] = critical_kb

    critical_lw = {}
    for cat, items in CRITICAL_ITEMS_LW.items():
        flagged = [i for i in items if i in r and r[i] is True]
        if flagged: critical_lw[cat] = flagged
    scores["critical_lw"] = critical_lw

    clinical_ts = [scores[f"{s}_T"] for s in ["Hs","D","Hy","Pd","Mf","Pa","Pt","Sc","Ma","Si"]]
    elevated = [t for t in clinical_ts if t > 44]
    scores["profile_elevation"] = round(sum(elevated)/len(elevated), 1) if elevated else 50.0

    scale_order = [("Hs",1),("D",2),("Hy",3),("Pd",4),("Mf",5),
                   ("Pa",6),("Pt",7),("Sc",8),("Ma",9),("Si",0)]
    sorted_scales = sorted(scale_order, key=lambda x: scores[f"{x[0]}_T"], reverse=True)
    def elev_sym(t):
        if t>=120: return "!!"
        elif t>=110: return "!"
        elif t>=100: return "**"
        elif t>=90:  return "*"
        elif t>=80:  return '"'
        elif t>=70:  return "'"
        elif t>=60:  return "-"
        elif t>=50:  return "/"
        elif t>=40:  return ":"
        else:        return "#"
    prev_t = None; welsh = ""
    for s, n in sorted_scales:
        t = scores[f"{s}_T"]
        if prev_t is not None and prev_t != t:
            welsh += elev_sym(prev_t)
        welsh += str(n); prev_t = t
    welsh += elev_sym(prev_t)
    scores["welsh_code"] = welsh
    scores["high_point_pair"] = f"{sorted_scales[0][1]}-{sorted_scales[1][1]}"
    return scores

def check_validity(scores):
    flags = []; valid = True

    # ── Cannot Say ──────────────────────────────────────────────
    if scores["cannot_say"] >= 30:
        flags.append("INVALID: Cannot Say ≥ 30 items omitted. Protocol should not be interpreted.")
        valid = False
    elif scores["cannot_say"] >= 10:
        flags.append(f"CAUTION: {scores['cannot_say']} items omitted (Cannot Say). Interpret with caution. "
                     f"If omissions are after item 370, Clinical/Validity scales may be interpreted "
                     f"but NOT Content, Supplementary, or Harris-Lingoes Scales.")

    # ── VRIN ────────────────────────────────────────────────────
    vrin_t = scores["VRIN_T"]
    if vrin_t >= 80:
        flags.append(f"INVALID: VRIN T={vrin_t} — inconsistent random responding. Protocol invalid.")
        valid = False
    elif vrin_t >= 70:
        flags.append(f"CAUTION: VRIN T={vrin_t} — possible inconsistent responding. Interpret with caution.")

    # ── TRIN ────────────────────────────────────────────────────
    trin_t = scores["TRIN_T"]
    if trin_t >= 80:
        flags.append(f"CAUTION: TRIN T={trin_t} — acquiescence or counter-acquiescence response set detected.")
    elif trin_t >= 70:
        flags.append(f"NOTE: TRIN T={trin_t} — suspect response set; possible uniform True or False responding.")

    # ── F Scale ─────────────────────────────────────────────────
    f_t = scores["F_T"]
    if f_t >= 100:
        flags.append(f"CAUTION: F T={f_t} — extremely elevated. Consider random responding, faking bad, "
                     f"or acute severe psychopathology. Protocol may be invalid unless VRIN is acceptable "
                     f"and client is a confirmed psychiatric inpatient.")
    elif f_t >= 80:
        flags.append(f"CAUTION: F T={f_t} — elevated. Possible symptom exaggeration, plea for help, "
                     f"or genuine severe psychopathology. Interpret clinical scales cautiously.")
    elif f_t >= 65:
        flags.append(f"NOTE: F T={f_t} — mildly elevated. May reflect genuine distress or specific problem areas.")

    # ── Fb Scale ────────────────────────────────────────────────
    fb_t = scores["Fb_T"]
    if fb_t >= 110:
        flags.append(f"INVALID: Fb T={fb_t} — do NOT interpret Content or Harris-Lingoes Scales.")
    elif fb_t >= 89:
        if f_t >= 89:
            flags.append(f"CAUTION: Fb T={fb_t} — do NOT interpret Content or Harris-Lingoes Scales.")
        else:
            flags.append(f"NOTE: Fb T={fb_t} — elevated back-half infrequency. Consider exaggeration or fatigue.")

    # ── Fp Scale ────────────────────────────────────────────────
    fp_t = scores["Fp_T"]
    if fp_t >= 100:
        flags.append(f"CRITICAL: Fp T={fp_t} — items rarely endorsed even by confirmed psychiatric inpatients. "
                     f"Strongly suggests symptom fabrication or gross exaggeration. "
                     f"Malingering evaluation is indicated. Profile validity is seriously compromised.")
        valid = False
    elif fp_t >= 80:
        flags.append(f"CAUTION: Fp T={fp_t} — significantly elevated. Possible malingering or dramatic symptom overreporting. "
                     f"Interpret clinical profile with significant caution.")
    elif fp_t >= 70:
        flags.append(f"NOTE: Fp T={fp_t} — mildly elevated. Consider alongside F and VRIN.")

    # ── L Scale ─────────────────────────────────────────────────
    l_t = scores["L_T"]
    if l_t >= 65:
        flags.append(f"NOTE: L T={l_t} — overly virtuous presentation; possible faking good. Interpret cautiously.")
    elif l_t >= 60:
        flags.append(f"NOTE: L T={l_t} — mildly elevated; tendency toward denial and conformity.")

    # ── K Scale ─────────────────────────────────────────────────
    k_t = scores["K_T"]
    if k_t >= 65:
        flags.append(f"NOTE: K T={k_t} — defensive responding; possible underreporting of psychological distress.")
    elif k_t <= 40:
        flags.append(f"NOTE: K T={k_t} — low defensiveness; openly acknowledging significant distress.")

    # ── S Scale ─────────────────────────────────────────────────
    s_t = scores["S_T"]
    if s_t <= 38:
        flags.append(f"NOTE: S T={s_t} — low Superlative Self-Presentation; not presenting in overly positive manner.")
    elif s_t >= 70:
        flags.append(f"NOTE: S T={s_t} — elevated Superlative Self-Presentation; consider defensiveness.")

    # ── F-K Index ───────────────────────────────────────────────
    fk = scores["FK_index"]
    if fk > 12:
        flags.append(f"CAUTION: F-K = {fk} (>12) — strong indicator of faking bad.")
    elif fk > 9:
        flags.append(f"NOTE: F-K = {fk} (>9) — possible symptom exaggeration.")
    elif fk < -12:
        flags.append(f"CAUTION: F-K = {fk} (<-12) — strong indicator of faking good.")
    elif fk < -9:
        flags.append(f"NOTE: F-K = {fk} (<-9) — possible symptom minimization.")

    # ── Critical Items — Suicidal Ideation Risk Alert ────────────
    critical_kb = scores.get("critical_kb", {})
    suicidal_items = critical_kb.get("Depressed Suicidal Ideation", [])
    if len(suicidal_items) >= 3:
        flags.append(f"⚠ RISK ALERT: {len(suicidal_items)} Koss-Butcher Depressed Suicidal Ideation critical items endorsed "
                     f"(Items: {', '.join(str(i) for i in suicidal_items)}). "
                     f"Immediate clinical risk assessment is required.")
    elif len(suicidal_items) >= 1:
        flags.append(f"NOTE: {len(suicidal_items)} Koss-Butcher Depressed Suicidal Ideation item(s) endorsed. "
                     f"Clinical follow-up regarding suicidal ideation is recommended.")

    threatened_assault = critical_kb.get("Threatened Assault", [])
    if threatened_assault:
        flags.append(f"⚠ RISK ALERT: Threatened Assault critical items endorsed "
                     f"(Items: {', '.join(str(i) for i in threatened_assault)}). "
                     f"Dangerousness assessment is indicated.")

    return {"valid": valid, "flags": flags}


# ══════════════════════════════════════════════════════════════
#  GROQ REPORT
# ══════════════════════════════════════════════════════════════

def generate_report(client_name, age, gender, scores, validity,
                    dob="Not provided", nationality="Not provided",
                    referral="Not provided"):
    clinical = {
        "Hs (Scale 1)": scores["Hs_T"], "D (Scale 2)": scores["D_T"],
        "Hy (Scale 3)": scores["Hy_T"], "Pd (Scale 4)": scores["Pd_T"],
        "Mf (Scale 5)": scores["Mf_T"], "Pa (Scale 6)": scores["Pa_T"],
        "Pt (Scale 7)": scores["Pt_T"], "Sc (Scale 8)": scores["Sc_T"],
        "Ma (Scale 9)": scores["Ma_T"], "Si (Scale 0)": scores["Si_T"],
    }
    content = {s: scores[f"{s}_T"] for s in
               ["ANX","FRS","OBS","DEP","HEA","BIZ","ANG","CYN","ASP","TPA","LSE","SOD","FAM","WRK","TRT"]}
    psy5 = {s: scores[f"{s}_T"] for s in ["AGGR","PSYC","DISC","NEGE","INTR"]}
    supp = {s: scores[f"{s}_T"] for s in ["A","R","MAC_R","Es","Do","Re","Mt","OH","APS","AAS","PK","Ho"]}
    elevated_content = {k:v for k,v in content.items() if v >= 65}
    validity_summary = "\n".join(validity["flags"]) if validity["flags"] else "No significant validity concerns."

    # Build critical items summary for prompt
    critical_kb_summary = ""
    if scores.get("critical_kb"):
        for cat, items in scores["critical_kb"].items():
            critical_kb_summary += f"  Koss-Butcher {cat}: Items {', '.join(str(i) for i in items)}\n"
    critical_lw_summary = ""
    if scores.get("critical_lw"):
        for cat, items in scores["critical_lw"].items():
            critical_lw_summary += f"  Lachar-Wrobel {cat}: Items {', '.join(str(i) for i in items)}\n"

    suicidal_count = len(scores.get("critical_kb", {}).get("Depressed Suicidal Ideation", []))
    assault_count  = len(scores.get("critical_kb", {}).get("Threatened Assault", []))

    prompt = f"""You are a licensed clinical psychologist writing a confidential MMPI-2 assessment report in English.
You must follow the interpretation rules below EXACTLY. These rules are authoritative and override any general knowledge.
Write the entire report in formal English. Do not use any Arabic words or phrases anywhere in the report.

═══════════════════════════════════════════════════════
MANDATORY INTERPRETATION RULES (follow precisely)
═══════════════════════════════════════════════════════

VALIDITY SCALES — REQUIRED INTERPRETATIONS:
- Cannot Say ≥30: Protocol INVALID, do not interpret
- Cannot Say 10–29: Interpret with caution; if omissions are after item 370, Clinical/Validity scales may be interpreted but NOT Content, Supplementary, or Harris-Lingoes
- VRIN ≥80: Protocol INVALID (random responding)
- VRIN 70–79: Possible invalid protocol; interpret cautiously
- TRIN ≥80: Response set detected (all-True or all-False tendency)
- F ≥100: May be invalid — consider random responding, faking bad, or acute psychosis
- F 80–99: Possible malingering OR genuine severe psychopathology OR plea for help
- F 65–79: May reflect deviant attitudes, genuine distress, or specific problem endorsement
- F <50: Socially conforming responses; may have faked good
- Fb ≥110: Invalid — do NOT interpret Content or Harris-Lingoes
- Fb ≥89 (when F ≥89): Do NOT interpret Content or Harris-Lingoes
- Fp ≥80: CRITICAL — items rarely endorsed even by confirmed psychiatric inpatients. Strongly suggests fabrication or gross exaggeration. State this prominently and explicitly.
- Fp 70–79: Some overreporting likely; consider alongside F and VRIN
- L ≥65: Overly virtuous presentation; profiles with L ≥65 should be interpreted very cautiously
- L 60–64: More conforming than average; tendency toward denial
- L <45: Responded frankly; admits minor faults; may be cynical
- K ≥65: Defensive; underreporting psychological distress
- K 41–55: Healthy balance; well-adjusted
- K ≤40: Openly acknowledging severe distress; possible exaggeration
- S ≤38: Not presenting in superlative positive manner; open about problems
- S ≥70: Highly positive self-presentation
- F-K >12: Strong indicator of faking bad
- F-K <-12: Strong indicator of faking good

CLINICAL SCALES — T-SCORE THRESHOLDS AND MEANINGS:
Scale 1 (Hs): T≥65 = excessive bodily concerns; vague somatic complaints; pessimistic; narcissistic; resists psychological interpretations; poor therapy candidate
Scale 2 (D): T≥65 = depressive symptoms; blue/dysphoric; pessimistic; guilt; psychomotor retardation; introverted; good therapy candidate
Scale 3 (Hy): T≥65 = reacts to stress via physical symptoms; immature/childish; self-centered; lacks insight; rarely reports hallucinations or delusions
Scale 4 (Pd): T≥65 = difficulty incorporating social values; rebellious; impulsive; poor judgment; narcissistic; superficial relationships; poor therapy prognosis
Scale 5 (Mf): Males T≥65 = aesthetic/artistic interests; sensitive; androgynous; NOT indicative of homosexuality; Males T<40 = stereotypically masculine; aggressive; inflexible
Scale 6 (Pa): T≥70 = frankly paranoid; delusions; ideas of reference; T 60–70 = paranoid predisposition; suspicious; guarded; poor therapy prognosis
Scale 7 (Pt): T≥85 = agitated ruminations no longer controlling anxiety; T≥65 = obsessive/anxious/tense; introspective; insecure; good therapy motivation but slow progress
Scale 8 (Sc): T≥91 = acute situational stress, typically NOT schizophrenic; T 65–90 = possible thought disorder; schizoid; isolated; withdrawn; poor therapy prognosis
Scale 9 (Ma): T≥80 = manic episode manifestations; T 65–79 = overactive; unrealistic self-appraisal; low frustration tolerance; poor therapy prognosis; T<35 = likely depressed
Scale 0 (Si): T≥65 = socially introverted; insecure socially; shy/timid; overcontrolled; may be depressed; T<40 = sociable/extroverted; impulse control problems

HARRIS-LINGOES SUBSCALES (T≥65):
  D1: unhappy/blue; lacks energy; nervous; poor concentration; socially withdrawn
  D2: immobile/withdrawn; denies aggressive impulses
  D3: preoccupied with physical functioning; denies good health
  D4: lacks energy; tense; memory/judgment difficulties
  D5: ruminates/cries; may feel life not worthwhile; losing control of thoughts
  Hy1: socially extroverted; comfortable interacting
  Hy2: naively optimistic/trusting; denies negative feelings
  Hy3: feels weak/fatigued; difficulties concentrating; unhappy
  Hy4: multiple somatic complaints; denies hostility
  Hy5: denies hostile/aggressive impulses; sensitive to others
  Pd1: unpleasant home/family; lacks support; feels controlled
  Pd2: resents parental/societal standards; has been in trouble
  Pd3: comfortable/confident socially; strong opinions
  Pd4: alienated/isolated; feels misunderstood; blames others
  Pd5: uncomfortable/unhappy; regret/guilt for past; may use alcohol
  Pa1: world as threatening; misunderstood; suspicious; may have persecution delusions
  Pa2: high-strung/sensitive; feels more intensely; lonely; seeks risky activities
  Pa3: extremely naive/optimistic; high moral standards; denies hostility
  Sc1: mistreated/unloved; others trying to harm; lonely/empty
  Sc2: fear/depression/apathy; may wish to be dead
  Sc3: fears losing mind; strange thoughts; unreality; concentration difficulties
  Sc4: life is a strain; despair; withdraws into fantasy; given up hope
  Sc5: not in control of emotions; restless/hyperactive; uncontrolled episodes
  Sc6: body changing strangely; hallucinations; ideas of external influence
  Ma1: perceives others as selfish/dishonest; feels justified acting similarly
  Ma2: accelerated speech/thought/activity; tense/restless; seeks excitement
  Ma3: denies social anxiety; impatient/irritable
  Ma4: unrealistic self-evaluation; resentful when others make demands
  Si1: shy; easily embarrassed; uncomfortable in new situations
  Si2: dislikes and avoids groups/crowds
  Si3: low self-esteem; self-critical; nervous/fearful; suspicious

CONTENT SCALES (T≥65):
  ANX: nervous/worried; concentration problems; sleep disturbance; pessimistic; overwhelmed
  FRS: fearful/uneasy; specific fears and phobias
  OBS: great difficulty making decisions; rigid; compulsive behaviors; intrusive thoughts
  DEP: depressed/despondent; fatigue; hopeless; preoccupied with death or suicide; indecisive; guilty
  HEA: denies good health; preoccupied with bodily functioning; multiple somatic symptoms
  BIZ: psychotic thought processes; hallucinations; feelings of unreality; beliefs others control mind
  ANG: angry/hostile; irritable/stubborn; temper tantrums; may be physically abusive
  CYN: others seen as dishonest/selfish; suspicious; guarded; untrusting
  ASP: trouble with law; antisocial behaviors; cynical; resents authority
  TPA: hard-driving; never enough time; hostile/irritable; holds grudges
  LSE: poor self-concept; anticipates failure; oversensitive to criticism; passive
  SOD: shy and socially introverted; prefers being alone; dislikes parties
  FAM: considerable family discord; families lacking love/support; angry toward family
  WRK: wide variety of behaviors contributing to poor work performance — NOTE: T≥90 indicates SEVERE work functioning impairment, not merely dissatisfaction
  TRT: negative attitudes toward mental health professionals; gives up easily; unable to make changes

SUPPLEMENTARY SCALES:
  Es T≥65 = stable/reliable; good therapy prognosis; Es T<40 = poor self-esteem; poor therapy prognosis
  MAC-R T≥65 = addictive personality; substance abuse likelihood
  A T≥65 = anxious; slow tempo; pessimistic; inhibited; motivated for therapy once trust established
  R T≥58 = submissive; conventional; avoids unpleasantness; unwilling to discuss problems
  Mt T≥65 = ineffectual; pessimistic; anxious and worried; procrastinates; somatic complaints; feels life is a strain
  IMPORTANT: Mt measures GENERAL MALADJUSTMENT AND DISTRESS — NEVER describe Mt as indicating manic or hypomanic symptoms
  OH T≥65 = fewer angry feelings; socialized; appropriate most of time but occasional exaggerated aggressive responses
  Do T≥65 = poised/self-assured; leadership; Do T<40 = submissive; lacks confidence
  Re T≥65 = strong moral sense; less likely criminal; Re T<40 = unwilling to accept responsibility; low integrity
  PK T≥65 = intense emotional distress; sleep disturbances; guilt/depression; unwanted disturbing thoughts; feels misunderstood
  APS/AAS T≥65 = substance abuse likelihood; AAS absence does NOT rule out substance abuse
  MDS T≥65 = significant marital problems
  Ho T≥65 = chronic hostility; cynical; suspicious
  AGGR T≥65 = offensive aggression; intimidates others; desire for power; likely violent
  PSYC T≥65 = disconnection from reality; delusions; disorganized thinking; perceptual aberrations
  DISC T≥65 = risk-taking; impulsive; rule-breaking; higher criminality
  NEGE T≥65 = worries; self-critical; guilty; nervous; negative affect
  INTR T≥65 = little joy or positive engagement; INTR T<40 = seeks social experiences; energetic

HIGH-POINT PAIRS (use when both scales T≥65):
  2-7/7-2: Worry, depression, anxiety, guilt; perfectionistic; excellent prognosis if distress manageable; pharmacotherapy if either >80; depressive or anxiety disorder diagnosis
  7-4/4-7: Cyclical acting out followed by guilt; insufficient controls; possible substance abuse; irritated by rules; guarded prognosis
  1-2/2-1: Depression with somatic complaints; passive-dependent; cynical about treatment; seeks medical attention
  2-8/8-2: Suicidal ideation likely; self-destructive potential high; major depression/schizophrenia/schizoaffective most likely; prognosis poor
  4-8/8-4: Irritable/hostile/suspicious; schizoid; suicidal attempts relatively common; borderline/schizoid/schizophrenia diagnoses
  7-8/8-7: Chronic worry; passivity; dependency; insecurity; sexual identity concerns; formal thought disorder possible; poor prognosis

CRITICAL ITEMS — MANDATORY REPORTING RULES:
- If Koss-Butcher Depressed Suicidal Ideation items are endorsed: YOU MUST explicitly address suicide risk in the report. State the number of items endorsed and that clinical risk assessment is required.
- If Threatened Assault items are endorsed: YOU MUST address dangerousness considerations.

═══════════════════════════════════════════════════════
CLIENT DATA
═══════════════════════════════════════════════════════

CLIENT: {client_name}
DATE OF BIRTH: {dob}
AGE: {age}
GENDER: {gender}
NATIONALITY: {nationality}
REFERRAL SOURCE: {referral}
TEST LANGUAGE: Arabic
ASSESSMENT: Minnesota Multiphasic Personality Inventory-2 (MMPI-2)
DATE: {datetime.datetime.now().strftime("%B %d, %Y")}
REPORT PREPARED BY: {THERAPIST_NAME}, {THERAPIST_TITLE} — {CENTER_NAME}

VALIDITY SCALE SUMMARY:
VRIN T={scores["VRIN_T"]} | TRIN T={scores["TRIN_T"]} | F T={scores["F_T"]} | Fb T={scores["Fb_T"]} | Fp T={scores["Fp_T"]}
L T={scores["L_T"]} | K T={scores["K_T"]} | S T={scores["S_T"]}
F-K Index: {scores["FK_index"]}
Validity Flags: {validity_summary}

CLINICAL SCALES (T-scores):
{chr(10).join(f"  Scale {k}: T={v}" for k,v in clinical.items())}

HIGH-POINT PAIR: {scores["high_point_pair"]}
WELSH CODE: {scores["welsh_code"]}
PROFILE ELEVATION: {scores["profile_elevation"]}

ELEVATED CONTENT SCALES (T≥65):
{chr(10).join(f"  {k}: T={v}" for k,v in elevated_content.items()) if elevated_content else "  None elevated"}

PSY-5 SCALES:
{chr(10).join(f"  {k}: T={v}" for k,v in psy5.items())}

SUPPLEMENTARY SCALES:
{chr(10).join(f"  {k}: T={v}" for k,v in supp.items())}

HARRIS-LINGOES ELEVATED (T≥65):
{chr(10).join(f"  {s}: T={scores[s+'_T']}" for s in ["D1","D2","D3","D4","D5","Hy1","Hy2","Hy3","Hy4","Hy5","Pd1","Pd2","Pd3","Pd4","Pd5","Pa1","Pa2","Pa3","Sc1","Sc2","Sc3","Sc4","Sc5","Sc6","Ma1","Ma2","Ma3","Ma4","Si1","Si2","Si3"] if scores[s+"_T"] >= 65) or "  None elevated"}

CRITICAL ITEMS ENDORSED:
{critical_kb_summary if critical_kb_summary else "  None flagged"}
{critical_lw_summary if critical_lw_summary else ""}
{"⚠ SUICIDAL IDEATION: " + str(suicidal_count) + " Koss-Butcher Depressed Suicidal Ideation items endorsed — MUST be addressed explicitly in report." if suicidal_count > 0 else ""}
{"⚠ THREATENED ASSAULT: Items endorsed — MUST be addressed explicitly in report." if assault_count > 0 else ""}

═══════════════════════════════════════════════════════
REPORT INSTRUCTIONS
═══════════════════════════════════════════════════════

Write a comprehensive professional MMPI-2 assessment report ENTIRELY IN ENGLISH with exactly these sections:

SECTION A — VALIDITY AND RESPONSE STYLE
Discuss all validity indicators using the rules above. State clearly whether the protocol is interpretable.
If Fp is elevated ≥70, this MUST be prominently discussed.
If suicidal ideation critical items are endorsed, state this here and require clinical follow-up.

SECTION B — CLINICAL SCALE ANALYSIS
Discuss the clinical profile using the interpretation rules above.
Address the high-point pair using the paired interpretation if both scales are ≥65.
Address ALL clinically elevated scales (T≥65) using the correct behavioral correlates.
Do NOT confuse scale names — Mt is maladjustment NOT mania.

SECTION C — HARRIS-LINGOES SUBSCALE ANALYSIS
Interpret only elevated Harris-Lingoes subscales (T≥65) using the descriptions above.

SECTION D — CONTENT SCALE ANALYSIS
Discuss all elevated content scales. WRK at T≥90 must be described as severe work functioning impairment.

SECTION E — PSY-5 AND SUPPLEMENTARY SCALES
Discuss elevated PSY-5 and key supplementary findings.
Mt MUST be described as general maladjustment/distress — NEVER as manic or hypomanic symptoms.

SECTION F — INTEGRATED CLINICAL FORMULATION
Synthesize all findings. If suicidal ideation critical items were endorsed, explicitly state risk assessment is required.

SECTION G — TREATMENT IMPLICATIONS
Evidence-based treatment recommendations. Note prognosis.

SECTION H — SUMMARY
One concise paragraph. Include any risk alerts prominently.

Use formal clinical English language. Reference specific T-scores throughout.
Label each section exactly as above. Do not reproduce the interpretation rules verbatim.
IMPORTANT: The entire report must be written in English only. Do not include any Arabic text."""

    api_key = st.secrets.get("GROQ_API_KEY", "")
    if not api_key:
        raise ValueError("GROQ_API_KEY is missing.")
    response = requests.post(
        "https://api.groq.com/openai/v1/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={"model": "llama-3.3-70b-versatile",
              "messages": [{"role": "user", "content": prompt}],
              "max_tokens": 3500, "temperature": 0.4},
        timeout=120,
    )
    if not response.ok:
        try:    ed = response.json()
        except: ed = response.text
        raise Exception(f"Groq error {response.status_code}: {ed}")
    return response.json()["choices"][0]["message"]["content"].strip()

# ══════════════════════════════════════════════════════════════
#  PROFILE CHART (Clinical Scales)
# ══════════════════════════════════════════════════════════════

def build_profile_chart(scores, width_pts=480, height_pts=160):
    """Build a clinical profile line chart as a ReportLab Drawing."""
    from reportlab.graphics.shapes import (
        Drawing, Line, Rect, String, PolyLine, Circle
    )
    from reportlab.lib.colors import HexColor, white, black

    scale_labels = ["Hs","D","Hy","Pd","Mf","Pa","Pt","Sc","Ma","Si"]
    scale_nums   = ["1","2","3","4","5","6","7","8","9","0"]
    t_values     = [scores[f"{s}_T"] for s in scale_labels]

    PAD_L, PAD_R, PAD_T, PAD_B = 30, 10, 10, 28
    chart_w = width_pts - PAD_L - PAD_R
    chart_h = height_pts - PAD_T - PAD_B
    T_MIN, T_MAX = 20, 120

    d = Drawing(width_pts, height_pts)

    # Background
    d.add(Rect(PAD_L, PAD_B, chart_w, chart_h,
               fillColor=HexColor("#FAFAFA"), strokeColor=HexColor("#CCCCCC"), strokeWidth=0.5))

    # Horizontal grid lines + T labels
    for t in [30,40,50,60,65,70,80,90,100,110,120]:
        y = PAD_B + (t - T_MIN) / (T_MAX - T_MIN) * chart_h
        col = HexColor("#FF9800") if t == 65 else HexColor("#DDDDDD")
        sw  = 0.8 if t == 65 else 0.4
        d.add(Line(PAD_L, y, PAD_L + chart_w, y, strokeColor=col, strokeWidth=sw))
        d.add(String(PAD_L - 4, y - 3, str(t),
                     fontSize=5, fillColor=HexColor("#888888"),
                     textAnchor="end"))

    # Vertical scale positions
    n = len(scale_labels)
    xs = [PAD_L + (i + 0.5) * chart_w / n for i in range(n)]

    # Vertical dividers
    for x in xs:
        d.add(Line(x, PAD_B, x, PAD_B + chart_h,
                   strokeColor=HexColor("#EEEEEE"), strokeWidth=0.3))

    # Scale labels at bottom
    for i, (lbl, num) in enumerate(zip(scale_labels, scale_nums)):
        d.add(String(xs[i], PAD_B - 10, num,
                     fontSize=6, fillColor=HexColor("#555555"), textAnchor="middle"))
        d.add(String(xs[i], PAD_B - 18, lbl,
                     fontSize=5, fillColor=HexColor("#888888"), textAnchor="middle"))

    # Line connecting points
    ys = [PAD_B + (t - T_MIN) / (T_MAX - T_MIN) * chart_h for t in t_values]
    for i in range(len(xs)-1):
        d.add(Line(xs[i], ys[i], xs[i+1], ys[i+1],
                   strokeColor=HexColor("#1A5CB8"), strokeWidth=1.2))

    # Data points
    for i, (x, y, t) in enumerate(zip(xs, ys, t_values)):
        dot_col = HexColor("#D9534F") if t >= 80 else HexColor("#F0AD4E") if t >= 65 else HexColor("#1A5CB8")
        d.add(Circle(x, y, 3, fillColor=dot_col, strokeColor=white, strokeWidth=0.5))
        # T value label above dot
        d.add(String(x, y + 5, str(t),
                     fontSize=5, fillColor=dot_col, textAnchor="middle"))

    return d

# ══════════════════════════════════════════════════════════════
#  PDF CREATION
# ══════════════════════════════════════════════════════════════

def create_pdf(path, client_name, age, gender, scores, validity, report_text,
               dob="Not provided", nationality="Not provided", referral="Not provided"):
    DARK   = colors.HexColor("#1C1917")
    WARM   = colors.HexColor("#6B5B45")
    LIGHT  = colors.HexColor("#F7F3EE")
    BORDER = colors.HexColor("#DDD5C8")
    RED    = colors.HexColor("#B71C1C")
    ORANGE = colors.HexColor("#E65100")
    BLUE   = colors.HexColor("#1A5CB8")

    def t_color(t):
        if t is None: return DARK
        if t >= 80:   return RED
        elif t >= 65: return ORANGE
        elif t <= 40: return BLUE
        return DARK

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    title_s  = ParagraphStyle("T",  fontName="Times-Roman",      fontSize=18, textColor=DARK, alignment=TA_CENTER, spaceAfter=15)
    sub_s    = ParagraphStyle("S",  fontName="Times-Italic",      fontSize=10, textColor=WARM, alignment=TA_CENTER, spaceAfter=2)
    meta_s   = ParagraphStyle("M",  fontName="Helvetica",         fontSize=8,  textColor=WARM, alignment=TA_CENTER, spaceAfter=10)
    sec_s    = ParagraphStyle("Se", fontName="Helvetica-Bold",    fontSize=10, textColor=WARM, spaceBefore=12, spaceAfter=4)
    body_s   = ParagraphStyle("B",  fontName="Helvetica",         fontSize=9.5,textColor=DARK, leading=15, spaceAfter=5)
    small_s  = ParagraphStyle("Sm", fontName="Helvetica",         fontSize=8,  textColor=WARM, leading=12)
    warn_s   = ParagraphStyle("W",  fontName="Helvetica-Bold",    fontSize=9,  textColor=RED,  leading=14, spaceAfter=4)
    footer_s = ParagraphStyle("F",  fontName="Helvetica-Oblique", fontSize=7,  textColor=WARM, leading=10, alignment=TA_CENTER)

    story = []
    date_str = datetime.datetime.now().strftime("%B %d, %Y  |  %H:%M")

    if os.path.exists(LOGO_FILE):
        try:
            logo = RLImage(LOGO_FILE, width=3.5*cm, height=1.8*cm)
            logo.hAlign = "CENTER"
            story.append(logo); story.append(Spacer(1, 0.2*cm))
        except: pass

    story += [
        Paragraph("Minnesota Multiphasic Personality Inventory-2", title_s),
        Paragraph("MMPI-2 — Extended Score Report", sub_s),
        Paragraph(f"CONFIDENTIAL  ·  {date_str}", meta_s),
        HRFlowable(width="100%", thickness=1, color=BORDER), Spacer(1, 0.3*cm),
    ]

    # Client info — expanded 4-row table
    info_data = [
        [Paragraph("<b>Client</b>", small_s), Paragraph(client_name, body_s),
         Paragraph("<b>Age</b>", small_s), Paragraph(str(age), body_s),
         Paragraph("<b>Gender</b>", small_s), Paragraph(gender, body_s)],
        [Paragraph("<b>Date of Birth</b>", small_s), Paragraph(dob, body_s),
         Paragraph("<b>Nationality</b>", small_s), Paragraph(nationality, body_s),
         Paragraph("<b>Test Language</b>", small_s), Paragraph("English", body_s)],
        [Paragraph("<b>Referral Source</b>", small_s), Paragraph(referral, body_s),
         Paragraph("<b>Assessment</b>", small_s), Paragraph("MMPI-2 (567 items)", body_s),
         Paragraph("<b>Date</b>", small_s), Paragraph(date_str, body_s)],
        [Paragraph("<b>Welsh Code</b>", small_s), Paragraph(scores["welsh_code"], body_s),
         Paragraph("<b>Prepared by</b>", small_s),
         Paragraph(f"{THERAPIST_NAME}, {THERAPIST_TITLE}", body_s),
         Paragraph("<b>Center</b>", small_s), Paragraph(CENTER_NAME, body_s)],
    ]
    it = Table(info_data, colWidths=[2.8*cm, 4.2*cm, 2.5*cm, 3.2*cm, 2.5*cm, 2.0*cm])
    it.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,-1),LIGHT), ("BOX",(0,0),(-1,-1),0.5,BORDER),
        ("INNERGRID",(0,0),(-1,-1),0.3,BORDER),
        ("TOPPADDING",(0,0),(-1,-1),6), ("BOTTOMPADDING",(0,0),(-1,-1),6),
        ("LEFTPADDING",(0,0),(-1,-1),8),
    ]))
    story += [it, Spacer(1, 0.3*cm)]

    # Validity flags
    if validity["flags"]:
        story.append(Paragraph("VALIDITY FLAGS", sec_s))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
        for flag in validity["flags"]:
            story.append(Paragraph(f"• {flag}", warn_s))
        story.append(Spacer(1, 0.2*cm))

    # Validity scales
    story.append(Paragraph("VALIDITY SCALES", sec_s))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 0.2*cm))

    def bar_str(t, width=20):
        if t is None: t = 50
        filled = max(0, min(width, int(((t-20)/100)*width)))
        hex_c = "#D9534F" if t>=80 else "#F0AD4E" if t>=65 else "#4A90D9" if t<=40 else "#4CAF50"
        return f'<font color="{hex_c}">{"█"*filled}</font><font color="#CCCCCC">{"░"*(width-filled)}</font>'

    val_rows = [[Paragraph("<b>Scale</b>",small_s), Paragraph("<b>Raw</b>",small_s),
                 Paragraph("<b>T</b>",small_s), Paragraph("<b>Profile</b>",small_s)]]
    for name, rk, tk in [
        ("VRIN","VRIN_raw","VRIN_T"),("TRIN","TRIN_raw","TRIN_T"),
        ("F","F_raw","F_T"),("Fb","Fb_raw","Fb_T"),("Fp","Fp_raw","Fp_T"),
        ("L","L_raw","L_T"),("K","K_raw","K_T"),("S","S_raw","S_T"),
    ]:
        t = scores.get(tk) or 50; raw = scores.get(rk) or 0
        val_rows.append([
            Paragraph(name, small_s),
            Paragraph(str(raw), ParagraphStyle("v",fontName="Helvetica",fontSize=8,alignment=TA_CENTER)),
            Paragraph(f"<b>{t}</b>", ParagraphStyle("vt",fontName="Helvetica-Bold",fontSize=9,textColor=t_color(t),alignment=TA_CENTER)),
            Paragraph(bar_str(t), ParagraphStyle("vb",fontName="Courier",fontSize=7)),
        ])
    vt = Table(val_rows, colWidths=[2.5*cm, 1.5*cm, 1.5*cm, 11.5*cm])
    vt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#EDE9E3")),
        ("BOX",(0,0),(-1,-1),0.5,BORDER), ("INNERGRID",(0,0),(-1,-1),0.3,BORDER),
        ("TOPPADDING",(0,0),(-1,-1),4), ("BOTTOMPADDING",(0,0),(-1,-1),4),
        ("LEFTPADDING",(0,0),(-1,-1),6), ("ALIGN",(1,0),(2,-1),"CENTER"),
    ]))
    story += [vt, Spacer(1, 0.4*cm)]

    # Clinical scales table
    story.append(Paragraph("CLINICAL SCALES", sec_s))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 0.2*cm))

    clin_data = [
        ("1 — Hs  Hypochondriasis",      "Hs_T","Hs_k"),
        ("2 — D   Depression",            "D_T", "D_raw"),
        ("3 — Hy  Hysteria",              "Hy_T","Hy_raw"),
        ("4 — Pd  Psychopathic Deviate",  "Pd_T","Pd_k"),
        ("5 — Mf  Masculinity-Femininity","Mf_T","Mf_raw"),
        ("6 — Pa  Paranoia",              "Pa_T","Pa_raw"),
        ("7 — Pt  Psychasthenia",         "Pt_T","Pt_k"),
        ("8 — Sc  Schizophrenia",         "Sc_T","Sc_k"),
        ("9 — Ma  Hypomania",             "Ma_T","Ma_k"),
        ("0 — Si  Social Introversion",   "Si_T","Si_raw"),
    ]
    clin_rows = [[Paragraph("<b>Scale</b>",small_s), Paragraph("<b>Raw</b>",small_s),
                  Paragraph("<b>T</b>",small_s), Paragraph("<b>Profile (20–120)</b>",small_s)]]
    for label, tk, rk in clin_data:
        t = scores.get(tk) or 50; raw = scores.get(rk) or 0
        clin_rows.append([
            Paragraph(label, small_s),
            Paragraph(str(raw), ParagraphStyle("c",fontName="Helvetica",fontSize=8,alignment=TA_CENTER)),
            Paragraph(f"<b>{t}</b>", ParagraphStyle("ct",fontName="Helvetica-Bold",fontSize=9,textColor=t_color(t),alignment=TA_CENTER)),
            Paragraph(bar_str(t), ParagraphStyle("cb",fontName="Courier",fontSize=7)),
        ])
    ct_table = Table(clin_rows, colWidths=[5.5*cm, 1.5*cm, 1.5*cm, 8.5*cm])
    ct_table.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#EDE9E3")),
        ("BOX",(0,0),(-1,-1),0.5,BORDER), ("INNERGRID",(0,0),(-1,-1),0.3,BORDER),
        ("TOPPADDING",(0,0),(-1,-1),4), ("BOTTOMPADDING",(0,0),(-1,-1),4),
        ("LEFTPADDING",(0,0),(-1,-1),6), ("ALIGN",(1,0),(2,-1),"CENTER"),
    ]))
    for i, (_, tk, _) in enumerate(clin_data, 1):
        if (scores.get(tk) or 0) >= 65:
            ct_table.setStyle(TableStyle([("BACKGROUND",(0,i),(-1,i),colors.HexColor("#FFF3F3"))]))
    story += [ct_table, Spacer(1, 0.2*cm)]

    # Profile elevation summary
    story.append(Paragraph(
        f"Profile Elevation: {scores['profile_elevation']}  |  "
        f"High-Point Pair: {scores['high_point_pair']}  |  "
        f"Welsh Code: {scores['welsh_code']}  |  F-K Index: {scores['FK_index']}",
        ParagraphStyle("pe", fontName="Helvetica", fontSize=8, textColor=WARM, leading=11)
    ))
    story.append(Spacer(1, 0.3*cm))

    # ── PROFILE CHART ──────────────────────────────────────────
    story.append(Paragraph("CLINICAL PROFILE CHART", sec_s))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 0.2*cm))
    chart_drawing = build_profile_chart(scores, width_pts=480, height_pts=170)
    story.append(chart_drawing)
    story.append(Paragraph(
        "<i>Profile chart: T-scores for Clinical Scales 1(Hs) through 0(Si). "
        "Orange line = T=65 clinical threshold. Red dots = T≥80; Orange = T 65–79; Blue = T≤40.</i>",
        ParagraphStyle("cn", fontName="Helvetica-Oblique", fontSize=7, textColor=WARM, leading=10, spaceAfter=4)
    ))
    story += [Spacer(1, 0.3*cm), PageBreak()]

    # Content scales
    story.append(Paragraph("CONTENT SCALES", sec_s))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 0.2*cm))
    cont_scales = ["ANX","FRS","OBS","DEP","HEA","BIZ","ANG","CYN","ASP","TPA","LSE","SOD","FAM","WRK","TRT"]
    cont_labels = {"ANX":"Anxiety","FRS":"Fears","OBS":"Obsessiveness","DEP":"Depression",
                   "HEA":"Health Concerns","BIZ":"Bizarre Mentation","ANG":"Anger",
                   "CYN":"Cynicism","ASP":"Antisocial Practices","TPA":"Type A",
                   "LSE":"Low Self-Esteem","SOD":"Social Discomfort","FAM":"Family Problems",
                   "WRK":"Work Interference","TRT":"Negative Treatment Indicators"}
    cont_rows = [[Paragraph("<b>Scale</b>",small_s), Paragraph("<b>Raw</b>",small_s),
                  Paragraph("<b>T</b>",small_s), Paragraph("<b>Profile</b>",small_s)]]
    for s in cont_scales:
        t = scores.get(f"{s}_T") or 50; raw = scores.get(f"{s}_raw") or 0
        cont_rows.append([
            Paragraph(f"{s} — {cont_labels[s]}", small_s),
            Paragraph(str(raw), ParagraphStyle("d",fontName="Helvetica",fontSize=8,alignment=TA_CENTER)),
            Paragraph(f"<b>{t}</b>", ParagraphStyle("dt",fontName="Helvetica-Bold",fontSize=9,textColor=t_color(t),alignment=TA_CENTER)),
            Paragraph(bar_str(t), ParagraphStyle("db",fontName="Courier",fontSize=7)),
        ])
    cont_t = Table(cont_rows, colWidths=[5.5*cm, 1.5*cm, 1.5*cm, 8.5*cm])
    cont_t.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#EDE9E3")),
        ("BOX",(0,0),(-1,-1),0.5,BORDER), ("INNERGRID",(0,0),(-1,-1),0.3,BORDER),
        ("TOPPADDING",(0,0),(-1,-1),4), ("BOTTOMPADDING",(0,0),(-1,-1),4),
        ("LEFTPADDING",(0,0),(-1,-1),6), ("ALIGN",(1,0),(2,-1),"CENTER"),
    ]))
    for i, s in enumerate(cont_scales, 1):
        if (scores.get(f"{s}_T") or 0) >= 65:
            cont_t.setStyle(TableStyle([("BACKGROUND",(0,i),(-1,i),colors.HexColor("#FFF3F3"))]))
    story += [cont_t, Spacer(1, 0.4*cm)]

    # Supplementary + PSY-5
    story.append(Paragraph("SUPPLEMENTARY & PSY-5 SCALES", sec_s))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 0.2*cm))
    supp_list = [("A","Anxiety"),("R","Repression"),("MAC_R","MacAndrew Alcoholism-R"),
                 ("Es","Ego Strength"),("Do","Dominance"),("Re","Social Responsibility"),
                 ("Mt","College Maladjustment"),("OH","Overcontrolled Hostility"),
                 ("APS","Addiction Potential"),("AAS","Addiction Acknowledgment"),
                 ("MDS","Marital Distress"),("Ho","Cook-Medley Hostility"),("PK","PTSD-Keane"),
                 ("GM","Masculine Gender Role"),("GF","Feminine Gender Role")]
    psy5_list = [("AGGR","Aggressiveness"),("PSYC","Psychoticism"),("DISC","Disconstraint"),
                 ("NEGE","Neg Emotionality/Neuroticism"),("INTR","Introversion/Low Pos Emotion")]

    def small_tbl(rows, widths):
        t = Table(rows, colWidths=widths)
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#EDE9E3")),
            ("BOX",(0,0),(-1,-1),0.5,BORDER), ("INNERGRID",(0,0),(-1,-1),0.3,BORDER),
            ("TOPPADDING",(0,0),(-1,-1),3), ("BOTTOMPADDING",(0,0),(-1,-1),3),
            ("LEFTPADDING",(0,0),(-1,-1),5), ("ALIGN",(1,0),(2,-1),"CENTER"),
        ]))
        return t

    sr = [[Paragraph("<b>Scale</b>",small_s),Paragraph("<b>Raw</b>",small_s),Paragraph("<b>T</b>",small_s)]]
    for s, lbl in supp_list:
        t = scores.get(f"{s}_T") or 50; raw = scores.get(f"{s}_raw") or 0
        sr.append([Paragraph(f"{s} — {lbl}",small_s),
                   Paragraph(str(raw),ParagraphStyle("sr",fontName="Helvetica",fontSize=8,alignment=TA_CENTER)),
                   Paragraph(f"<b>{t}</b>",ParagraphStyle("st",fontName="Helvetica-Bold",fontSize=9,textColor=t_color(t),alignment=TA_CENTER))])
    pr = [[Paragraph("<b>PSY-5 Scale</b>",small_s),Paragraph("<b>Raw</b>",small_s),Paragraph("<b>T</b>",small_s)]]
    for s, lbl in psy5_list:
        t = scores.get(f"{s}_T") or 50; raw = scores.get(f"{s}_raw") or 0
        pr.append([Paragraph(f"{s} — {lbl}",small_s),
                   Paragraph(str(raw),ParagraphStyle("pr2",fontName="Helvetica",fontSize=8,alignment=TA_CENTER)),
                   Paragraph(f"<b>{t}</b>",ParagraphStyle("pt2",fontName="Helvetica-Bold",fontSize=9,textColor=t_color(t),alignment=TA_CENTER))])

    story += [small_tbl(sr,[5*cm,1.2*cm,1.2*cm]), Spacer(1,0.4*cm),
              small_tbl(pr,[5.5*cm,1.2*cm,1.2*cm]), Spacer(1,0.4*cm), PageBreak()]

    # Harris-Lingoes
    story.append(Paragraph("HARRIS-LINGOES SUBSCALES", sec_s))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 0.2*cm))
    hl_groups = {
        "Depression":           [("D1","Subjective Depression"),("D2","Psychomotor Retardation"),("D3","Physical Malfunctioning"),("D4","Mental Dullness"),("D5","Brooding")],
        "Hysteria":             [("Hy1","Denial of Social Anxiety"),("Hy2","Need for Affection"),("Hy3","Lassitude-Malaise"),("Hy4","Somatic Complaints"),("Hy5","Inhibition of Aggression")],
        "Psychopathic Deviate": [("Pd1","Familial Discord"),("Pd2","Authority Problems"),("Pd3","Social Imperturbability"),("Pd4","Social Alienation"),("Pd5","Self-Alienation")],
        "Paranoia":             [("Pa1","Persecutory Ideas"),("Pa2","Poignancy"),("Pa3","Naivete")],
        "Schizophrenia":        [("Sc1","Social Alienation"),("Sc2","Emotional Alienation"),("Sc3","Lack of Ego Mastery-Cognitive"),("Sc4","Lack of Ego Mastery-Conative"),("Sc5","Lack of Ego Mastery-Defective Inhibition"),("Sc6","Bizarre Sensory Experiences")],
        "Hypomania":            [("Ma1","Amorality"),("Ma2","Psychomotor Acceleration"),("Ma3","Imperturbability"),("Ma4","Ego Inflation")],
        "Social Introversion":  [("Si1","Shyness/Self-Consciousness"),("Si2","Social Avoidance"),("Si3","Alienation-Self and Others")],
    }
    hl_rows = [[Paragraph("<b>Subscale</b>",small_s),Paragraph("<b>Raw</b>",small_s),Paragraph("<b>T</b>",small_s)]]
    for group, items in hl_groups.items():
        hl_rows.append([Paragraph(f"<b>{group}</b>",ParagraphStyle("gg",fontName="Helvetica-Bold",fontSize=8,textColor=WARM)),"",""])
        for code, lbl in items:
            t = scores.get(f"{code}_T") or 50; raw = scores.get(f"{code}_raw") or 0
            hl_rows.append([
                Paragraph(f"  {code} — {lbl}", small_s),
                Paragraph(str(raw),ParagraphStyle("hr",fontName="Helvetica",fontSize=8,alignment=TA_CENTER)),
                Paragraph(f"<b>{t}</b>",ParagraphStyle("ht",fontName="Helvetica-Bold",fontSize=9,textColor=t_color(t),alignment=TA_CENTER)),
            ])
    hl_t = Table(hl_rows, colWidths=[10*cm,1.5*cm,1.5*cm])
    hl_t.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#EDE9E3")),
        ("BOX",(0,0),(-1,-1),0.5,BORDER), ("INNERGRID",(0,0),(-1,-1),0.3,BORDER),
        ("TOPPADDING",(0,0),(-1,-1),3), ("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),6), ("ALIGN",(1,0),(2,-1),"CENTER"),
    ]))
    story += [hl_t, Spacer(1,0.4*cm), PageBreak()]

    # Critical items
    if scores["critical_kb"] or scores["critical_lw"]:
        story.append(Paragraph("CRITICAL ITEMS", sec_s))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
        story.append(Spacer(1, 0.2*cm))
        story.append(Paragraph("Items endorsed in clinically significant direction. Single-item responses are unreliable — use for hypothesis generation only.",
            ParagraphStyle("ci",fontName="Helvetica-Oblique",fontSize=8,textColor=WARM,leading=12,spaceAfter=6)))
        for cat, items in scores["critical_kb"].items():
            story.append(Paragraph(f"<b>Koss-Butcher — {cat}:</b>  Items {', '.join(str(i) for i in items)}", body_s))
        for cat, items in scores["critical_lw"].items():
            story.append(Paragraph(f"<b>Lachar-Wrobel — {cat}:</b>  Items {', '.join(str(i) for i in items)}", body_s))
        story += [Spacer(1,0.3*cm), PageBreak()]

    # Clinical report
    story.append(Paragraph("CLINICAL REPORT", sec_s))
    story.append(HRFlowable(width="100%", thickness=1, color=BORDER))
    story.append(Spacer(1, 0.2*cm))
    for line in report_text.split("\n"):
        line = line.strip()
        if not line: story.append(Spacer(1, 0.15*cm))
        elif line.startswith("SECTION") or (line.isupper() and len(line) < 60):
            story.append(Paragraph(line, sec_s))
            story.append(HRFlowable(width="100%", thickness=0.4, color=BORDER))
        else:
            story.append(Paragraph(line, body_s))

    story += [
        Spacer(1, 0.5*cm),
        HRFlowable(width="100%", thickness=0.5, color=BORDER),
        Spacer(1, 0.3*cm),
    ]

    # Professional sign-off block
    sign_data = [[]]
    if os.path.exists(LOGO_FILE):
        try:
            sign_logo = RLImage(LOGO_FILE, width=2.2*cm, height=1.1*cm)
            sign_data = [[sign_logo,
                          Paragraph(
                              f"<b>{CENTER_NAME}</b><br/>"
                              f"{THERAPIST_NAME}<br/>"
                              f"<i>{THERAPIST_TITLE}</i>",
                              ParagraphStyle("sg", fontName="Helvetica", fontSize=8,
                                             textColor=WARM, leading=12)),
                          Paragraph(
                              "This report is strictly confidential and intended solely "
                              "for the treating clinician. Not to be shared without "
                              "explicit written consent.",
                              ParagraphStyle("fc", fontName="Helvetica-Oblique", fontSize=7,
                                             textColor=WARM, leading=10))]]
        except: pass

    if sign_data and sign_data[0]:
        sign_tbl = Table(sign_data, colWidths=[2.5*cm, 5.5*cm, 9.2*cm])
        sign_tbl.setStyle(TableStyle([
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
            ("LEFTPADDING",(0,0),(-1,-1),4),
            ("RIGHTPADDING",(0,0),(-1,-1),4),
            ("TOPPADDING",(0,0),(-1,-1),4),
            ("BOTTOMPADDING",(0,0),(-1,-1),4),
            ("LINEAFTER",(0,0),(0,0),0.5,BORDER),
            ("LINEAFTER",(1,0),(1,0),0.5,BORDER),
        ]))
        story.append(sign_tbl)
    else:
        story.append(Paragraph(
            f"{CENTER_NAME}  ·  {THERAPIST_NAME}, {THERAPIST_TITLE}  ·  Confidential",
            ParagraphStyle("fc2", fontName="Helvetica-Oblique", fontSize=7,
                           textColor=WARM, leading=10, alignment=TA_CENTER)))

    doc.build(story)

# ══════════════════════════════════════════════════════════════
#  ARABIC WORD DOCUMENT (Two-column table, compact)
# ══════════════════════════════════════════════════════════════

def set_cell_rtl(cell):
    """Set RTL direction on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    rtl = OxmlElement('w:textDirection')
    rtl.set(qn('w:val'), 'btLr')
    # Use paragraph RTL instead
    for para in cell.paragraphs:
        pPr = para._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.insert(0, bidi)

def set_para_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def create_arabic_word_doc(path, client_name, age, gender, responses):
    """Create compact two-column Arabic Q&A Word document."""
    doc = DocxDocument()

    # Page setup — narrow margins for density
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Title
    title = doc.add_paragraph()
    set_para_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("مقياس مينيسوتا متعدد الأوجه للشخصية - MMPI-2")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1C, 0x19, 0x17)

    subtitle = doc.add_paragraph()
    set_para_rtl(subtitle)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("استجابات المريض — للاستخدام في الجلسة العلاجية")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x6B, 0x5B, 0x45)

    # Client info line
    info_p = doc.add_paragraph()
    set_para_rtl(info_p)
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.datetime.now().strftime("%d/%m/%Y")
    info_run = info_p.add_run(f"الاسم: {client_name}     العمر: {age}     الجنس: {gender}     التاريخ: {date_str}")
    info_run.font.size = Pt(9)
    info_run.font.color.rgb = RGBColor(0x6B, 0x5B, 0x45)

    doc.add_paragraph()  # small gap

    # Two-column table: answer | question+number
    # Column widths: answer (narrow) | question (wide)
    # RTL layout: answer on left visually = col 0 in docx
    # But for Arabic reading: question right, answer left
    # Table: col0=number, col1=Arabic question, col2=answer — all 3 narrow
    # Actually: two visible columns but we need number inside question col

    # PAGE WIDTH content = 21 - 1.5*2 = 18cm
    # Col 0 = Answer: 2.5cm | Col 1 = Q# + Question: 15.5cm
    TBL_W  = int(18 * 914400 / 2.54)  # EMU
    COL_ANS = int(2.5 * 914400 / 2.54)
    COL_Q   = int(15.5 * 914400 / 2.54)

    # We'll use DXA (1440 = 1 inch = 2.54cm)
    # 18cm = 10205 DXA approx
    COL_ANS_DXA = 1418  # ~2.5cm
    COL_Q_DXA   = 8787  # ~15.5cm

    # Header row
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(15.5)

    hdr = tbl.rows[0]
    # Col 0: answer header
    c0 = hdr.cells[0]
    p0 = c0.paragraphs[0]
    set_para_rtl(p0)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("الإجابة")
    r0.bold = True; r0.font.size = Pt(9)
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c0._tc.get_or_add_tcPr().append(
        OxmlElement('w:shd') if False else _make_shading("2D2926")
    )

    # Col 1: question header
    c1 = hdr.cells[1]
    p1 = c1.paragraphs[0]
    set_para_rtl(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run("العبارة")
    r1.bold = True; r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c1._tc.get_or_add_tcPr().append(_make_shading("2D2926"))

    # Data rows — all 567 items
    for i, q_ar in enumerate(MMPI2_QUESTIONS_AR, start=1):
        ans = responses.get(i)
        if ans is True:
            ans_text = "صحيح ✓"
            ans_color = RGBColor(0x2E, 0x7D, 0x32)
        elif ans is False:
            ans_text = "خطأ ✗"
            ans_color = RGBColor(0xC6, 0x28, 0x28)
        else:
            ans_text = "—"
            ans_color = RGBColor(0x99, 0x99, 0x99)

        row = tbl.add_row()
        row.height = Cm(0.55)

        # Answer cell
        ca = row.cells[0]
        ca.width = Cm(2.5)
        pa = ca.paragraphs[0]
        set_para_rtl(pa)
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ra = pa.add_run(ans_text)
        ra.font.size = Pt(8)
        ra.bold = True
        ra.font.color.rgb = ans_color

        # Question cell
        cq = row.cells[1]
        cq.width = Cm(15.5)
        pq = cq.paragraphs[0]
        set_para_rtl(pq)
        pq.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Item number + question text
        rnum = pq.add_run(f"{i}. ")
        rnum.font.size = Pt(8)
        rnum.bold = True
        rnum.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
        rq = pq.add_run(q_ar)
        rq.font.size = Pt(8)
        rq.font.color.rgb = RGBColor(0x1C, 0x19, 0x17)

        # Zebra shading
        if i % 2 == 0:
            ca._tc.get_or_add_tcPr().append(_make_shading("F7F3EE"))
            cq._tc.get_or_add_tcPr().append(_make_shading("F7F3EE"))

    # Set compact row spacing on all paragraphs
    for row in tbl.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)

    doc.save(path)

def _make_shading(hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    return shd

# ══════════════════════════════════════════════════════════════
#  EMAIL
# ══════════════════════════════════════════════════════════════

def send_email(pdf_path, docx_path, client_name, scores, pdf_fname, docx_fname):
    date_str = datetime.datetime.now().strftime("%B %d, %Y at %H:%M")
    elevated = [(s, scores[f"{s}_T"]) for s in
                ["Hs","D","Hy","Pd","Mf","Pa","Pt","Sc","Ma","Si"]
                if scores[f"{s}_T"] >= 65]
    elev_html = "".join(
        f"<tr><td style='padding:4px 0;color:#6B5B45;'>Scale {s}</td><td><strong style='color:#D9534F;'>T={t}</strong></td></tr>"
        for s,t in elevated
    ) or "<tr><td colspan='2' style='color:#4CAF50;'>No clinical scales elevated ≥ 65</td></tr>"

    msg = MIMEMultipart("mixed")
    msg["From"] = GMAIL_ADDRESS; msg["To"] = THERAPIST_EMAIL
    msg["Subject"] = f"[MMPI-2 Report] {client_name} — {date_str}"
    body_html = f"""<html><body style="font-family:Georgia,serif;color:#1C1917;background:#F7F3EE;padding:24px;">
      <div style="max-width:580px;margin:0 auto;background:white;border:1px solid #DDD5C8;border-radius:4px;padding:32px;">
        <h2 style="font-weight:300;font-size:20px;margin-bottom:2px;">MMPI-2 Assessment Report</h2>
        <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
        <table style="width:100%;font-size:13px;border-collapse:collapse;">
          <tr><td style="padding:5px 0;color:#6B5B45;width:40%;">Client</td><td><strong>{client_name}</strong></td></tr>
          <tr><td style="padding:5px 0;color:#6B5B45;">Date</td><td>{date_str}</td></tr>
          <tr><td style="padding:5px 0;color:#6B5B45;">High-Point Pair</td><td><strong>{scores["high_point_pair"]}</strong></td></tr>
          <tr><td style="padding:5px 0;color:#6B5B45;">Profile Elevation</td><td>{scores["profile_elevation"]}</td></tr>
          <tr><td style="padding:5px 0;color:#6B5B45;">Welsh Code</td><td><code>{scores["welsh_code"]}</code></td></tr>
        </table>
        <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
        <p style="font-size:12px;color:#6B5B45;font-weight:bold;">Elevated Clinical Scales (T≥65)</p>
        <table style="width:100%;font-size:12px;border-collapse:collapse;">{elev_html}</table>
        <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
        <p style="font-size:12px;line-height:1.6;">Two files attached:<br>
        📄 <strong>English PDF</strong> — Full clinical report with profile chart<br>
        📝 <strong>Arabic Word Doc</strong> — Questions & answers for therapy session</p>
        <p style="font-size:10px;color:#6B5B45;font-style:italic;">Confidential — treating clinician only.</p>
      </div></body></html>"""
    msg.attach(MIMEText(body_html, "html"))
    for fpath, fname in [(pdf_path, pdf_fname), (docx_path, docx_fname)]:
        with open(fpath, "rb") as f:
            part = MIMEBase("application","octet-stream"); part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{fname}"')
        msg.attach(part)
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as srv:
        srv.login(GMAIL_ADDRESS, GMAIL_PASSWORD)
        srv.sendmail(GMAIL_ADDRESS, THERAPIST_EMAIL, msg.as_string())

# ══════════════════════════════════════════════════════════════
#  STREAMLIT UI — Arabic
# ══════════════════════════════════════════════════════════════

st.set_page_config(page_title="تقييم MMPI-2", page_icon="🧠",
                   layout="centered", initial_sidebar_state="collapsed")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@300;400;500&family=Jost:wght@300;400;500&display=swap');
:root{--cream:#F7F3EE;--deep:#1C1917;--warm:#8B7355;--accent:#C4956A;--border:#DDD5C8;--selected:#2D2926;}
#MainMenu{visibility:hidden!important;display:none!important;}
header[data-testid="stHeader"]{visibility:hidden!important;display:none!important;}
footer{visibility:hidden!important;display:none!important;}
[data-testid="stToolbar"],[data-testid="stDecoration"],[data-testid="stStatusWidget"],[data-testid="stActionButton"]{display:none!important;}
a[href*="streamlit.io"]{display:none!important;}
[class*="viewerBadge"],[class*="ProfileBadge"]{display:none!important;}
html,body,[data-theme="dark"],[data-theme="light"]{color-scheme:light only!important;}
[data-testid="stAppViewContainer"],.stApp{background-color:#F7F3EE!important;color:#1C1917!important;}
html,body,[class*="css"]{font-family:'Jost',sans-serif;background-color:var(--cream);color:var(--deep);direction:rtl;}
.stApp{background-color:var(--cream);}
.page-header{text-align:center;padding:2.5rem 0 1.5rem;border-bottom:1px solid var(--border);margin-bottom:1.5rem;direction:rtl;}
.page-header h1{font-family:'Cormorant Garamond',serif;font-size:2.2rem;font-weight:300;margin-bottom:.3rem;}
.page-header p{color:var(--warm);font-size:.82rem;letter-spacing:.05em;}
.q-card{background:white;border:1px solid var(--border);border-radius:4px;padding:1.2rem 1.5rem .5rem;margin-bottom:.8rem;direction:rtl;text-align:right;}
.q-num{font-size:.68rem;font-weight:500;letter-spacing:.06em;color:var(--accent);margin-bottom:.3rem;}
.q-text{font-family:'Cormorant Garamond',serif;font-size:1.05rem;color:var(--deep);line-height:1.6;margin-bottom:.8rem;}
div[data-testid="stRadio"]>label{display:none;}
div[data-testid="stRadio"]>div{gap:.4rem!important;flex-direction:row-reverse!important;flex-wrap:wrap!important;justify-content:flex-start!important;}
div[data-testid="stRadio"]>div>label{background:var(--cream)!important;border:1px solid var(--border)!important;border-radius:20px!important;padding:.4rem 1.2rem!important;cursor:pointer!important;font-size:.85rem!important;color:var(--deep)!important;font-family:'Jost',sans-serif!important;transition:all .15s ease!important;white-space:nowrap!important;}
div[data-testid="stRadio"]>div>label:hover{border-color:var(--accent)!important;background:#FDF9F4!important;}
.progress-wrap{background:var(--border);border-radius:2px;height:3px;margin:1rem 0 .5rem;}
.progress-fill{height:3px;border-radius:2px;background:linear-gradient(90deg,var(--warm),var(--accent));}
.stButton>button{background:var(--selected)!important;color:var(--cream)!important;border:none!important;padding:.75rem 2.5rem!important;font-family:'Jost',sans-serif!important;font-size:.82rem!important;font-weight:500!important;letter-spacing:.08em!important;border-radius:2px!important;transition:background .2s!important;}
.stButton>button:hover{background:var(--warm)!important;}
.thank-you{text-align:center;padding:5rem 2rem;direction:rtl;}
.thank-you h2{font-family:'Cormorant Garamond',serif;font-size:2.2rem;font-weight:300;margin-bottom:1rem;}
.thank-you p{color:var(--warm);font-size:.95rem;max-width:420px;margin:0 auto;line-height:1.9;}
div[data-testid="stTextInput"] input{background:white!important;border:1px solid var(--border)!important;border-radius:3px!important;font-family:'Jost',sans-serif!important;}
div[data-testid="stSelectbox"] div{background:white!important;border:1px solid var(--border)!important;border-radius:3px!important;}
</style>""", unsafe_allow_html=True)

page = st.query_params.get("page", "client")

# ── ADMIN ──────────────────────────────────────────────────────
if page == "admin":
    st.markdown('<div class="page-header"><p>بوابة المعالج</p><h1>تقارير MMPI-2</h1></div>', unsafe_allow_html=True)
    if "admin_auth" not in st.session_state: st.session_state.admin_auth = False
    if not st.session_state.admin_auth:
        pwd = st.text_input("كلمة المرور", type="password", placeholder="أدخل كلمة المرور")
        if st.button("دخول"):
            if pwd == st.secrets.get("ADMIN_PASSWORD",""):
                st.session_state.admin_auth = True; st.rerun()
            else: st.error("كلمة المرور غير صحيحة.")
    else:
        reports_dir = "reports"
        os.makedirs(reports_dir, exist_ok=True)
        files = sorted([f for f in os.listdir(reports_dir) if f.endswith((".pdf",".docx"))], reverse=True)
        if not files: st.info("لا توجد تقارير بعد.")
        else:
            st.markdown(f"**{len(files)} ملف محفوظ**")
            for fname in files:
                c1, c2 = st.columns([3,1])
                with c1: st.markdown(f"📄 `{fname}`")
                with c2:
                    mime = "application/pdf" if fname.endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    with open(os.path.join(reports_dir,fname),"rb") as f:
                        st.download_button("تحميل", data=f, file_name=fname, mime=mime, key=fname)
        if st.button("تسجيل الخروج"): st.session_state.admin_auth=False; st.rerun()

# ── CLIENT ──────────────────────────────────────────────────────
else:
    if "submitted"          not in st.session_state: st.session_state.submitted = False
    if "current_page"       not in st.session_state: st.session_state.current_page = 0
    if "responses"          not in st.session_state: st.session_state.responses = {}
    if "client_name_saved"  not in st.session_state: st.session_state.client_name_saved = ""
    if "age_saved"          not in st.session_state: st.session_state.age_saved = ""
    if "gender_saved"       not in st.session_state: st.session_state.gender_saved = "— اختر —"
    if "dob_saved"          not in st.session_state: st.session_state.dob_saved = ""
    if "nationality_saved"  not in st.session_state: st.session_state.nationality_saved = ""
    if "referral_saved"     not in st.session_state: st.session_state.referral_saved = ""
    if "access_granted"     not in st.session_state: st.session_state.access_granted = False

    # ── Access code gate ───────────────────────────────────────
    if not st.session_state.access_granted:
        if os.path.exists(LOGO_FILE):
            c1, c2, c3 = st.columns([1, 2, 1])
            with c2: st.image(LOGO_FILE, use_container_width=True)
        st.markdown("""<div class="page-header">
            <p>تقييم نفسي سري</p>
            <h1>MMPI-2</h1>
            <p>مقياس مينيسوتا متعدد الأوجه للشخصية</p>
        </div>""", unsafe_allow_html=True)
        st.markdown("""<div style="max-width:360px;margin:0 auto;padding:2rem 0;text-align:center;direction:rtl;">
            <p style="color:#8B7355;font-size:.9rem;margin-bottom:1.5rem;line-height:1.8;">
                هذا التقييم متاح للمرضى المحالين فقط.<br>
                يرجى إدخال رمز الوصول الذي زوّدك به معالجك.
            </p>
        </div>""", unsafe_allow_html=True)
        col_a, col_b, col_c = st.columns([1, 2, 1])
        with col_b:
            code = st.text_input("رمز الوصول", type="password",
                                 placeholder="أدخل رمز الوصول",
                                 label_visibility="collapsed")
            if st.button("دخول", use_container_width=True):
                valid_codes = [c.strip() for c in st.secrets.get("ACCESS_CODE", "").split(",")]
                if code.strip() in valid_codes:
                    st.session_state.access_granted = True
                    st.rerun()
                else:
                    st.markdown("""<div style="background:#FFF0F0;border-right:3px solid #D9534F;
                        border-left:none;padding:.8rem 1rem;border-radius:4px 0 0 4px;
                        font-size:.88rem;color:#7A1A1A;margin:.5rem 0;
                        direction:rtl;text-align:right;">
                        &#9888; رمز الوصول غير صحيح. يرجى المراجعة والمحاولة مرة أخرى.
                    </div>""", unsafe_allow_html=True)
        st.stop()

    total_pages = math.ceil(567 / ITEMS_PER_PAGE)

    if st.session_state.submitted:
        st.markdown("""<div class="thank-you">
            <h2>شكرًا لك</h2>
            <p>تم تسليم إجاباتك بنجاح.<br>سيتواصل معك المعالج في أقرب وقت.</p>
        </div>""", unsafe_allow_html=True)
        if st.session_state.get("email_error"):
            st.warning(f"تم حفظ التقرير لكن فشل الإرسال: {st.session_state.email_error}")
    else:
        if os.path.exists(LOGO_FILE):
            c1,c2,c3 = st.columns([1,2,1])
            with c2: st.image(LOGO_FILE, use_container_width=True)

        st.markdown("""<div class="page-header">
            <p>تقييم نفسي سري</p>
            <h1>MMPI-2</h1>
            <p>مقياس مينيسوتا متعدد الأوجه للشخصية</p>
        </div>""", unsafe_allow_html=True)

        cp = st.session_state.current_page

        # Client info on page 0
        if cp == 0:
            st.markdown("""<p style="font-size:.88rem;color:#8B7355;text-align:center;
                margin-bottom:1.5rem;line-height:1.9;direction:rtl;">
                يحتوي هذا الاستبيان على ٥٦٧ عبارة. لكل عبارة، حدد ما إذا كانت
                <strong>صحيحة</strong> أم <strong>خطأ</strong> بالنسبة لك.<br>
                أجب على جميع العبارات. إذا لم تكن متأكدًا، اختر ما ينطبق عليك
                <em>في أغلب الأحيان</em>.
            </p>""", unsafe_allow_html=True)

            col1, col2, col3 = st.columns(3)
            with col1:
                client_name = st.text_input(
                    "اسمك باللغة الإنجليزية (اختياري)",
                    placeholder="Your name in English",
                    value=st.session_state.client_name_saved,
                    key="client_name_input"
                )
                st.session_state.client_name_saved = client_name
                if any('\u0600' <= c <= '\u06ff' for c in (client_name or "")):
                    st.markdown('<div style="color:#D9534F;font-size:.82rem;direction:rtl;">⚠ يرجى كتابة اسمك باللغة الإنجليزية فقط.</div>', unsafe_allow_html=True)
                dob = st.text_input("تاريخ الميلاد", placeholder="DD/MM/YYYY",
                                    value=st.session_state.dob_saved, key="dob_input")
                st.session_state.dob_saved = dob
            with col2:
                age = st.text_input("العمر", placeholder="مثال: ٣٥",
                                    value=st.session_state.age_saved, key="age_input")
                st.session_state.age_saved = age
                nationality = st.text_input("الجنسية", placeholder="مثال: مصري",
                                            value=st.session_state.nationality_saved,
                                            key="nationality_input")
                st.session_state.nationality_saved = nationality
            with col3:
                gender_opts = ["— اختر —", "ذكر", "أنثى"]
                gender_idx  = gender_opts.index(st.session_state.gender_saved) if st.session_state.gender_saved in gender_opts else 0
                gender_ar   = st.selectbox("الجنس", gender_opts, index=gender_idx, key="gender_input")
                st.session_state.gender_saved = gender_ar
                referral = st.text_input("مصدر الإحالة", placeholder="مثال: د. أحمد حسن",
                                         value=st.session_state.referral_saved,
                                         key="referral_input")
                st.session_state.referral_saved = referral

        # Questions for current page
        start_idx = cp * ITEMS_PER_PAGE
        end_idx   = min(start_idx + ITEMS_PER_PAGE, 567)
        page_questions = list(enumerate(MMPI2_QUESTIONS_AR[start_idx:end_idx], start=start_idx+1))

        st.markdown("<br>", unsafe_allow_html=True)

        for item_num, q_text in page_questions:
            st.markdown(f"""<div class="q-card">
                <div class="q-num">عبارة {item_num} من ٥٦٧</div>
                <div class="q-text">{q_text}</div>
            </div>""", unsafe_allow_html=True)

            prev = st.session_state.responses.get(item_num)
            prev_idx = None
            if prev is True:  prev_idx = 0
            elif prev is False: prev_idx = 1

            choice = st.radio("", ["صحيح", "خطأ"], index=prev_idx,
                              key=f"item_{item_num}", horizontal=True,
                              label_visibility="collapsed")
            if choice is not None:
                st.session_state.responses[item_num] = (choice == "صحيح")

        # Progress
        answered = len(st.session_state.responses)
        pct = int((answered / 567) * 100)
        st.markdown(f"""<div style="text-align:center;font-size:.78rem;color:#8B7355;
                        letter-spacing:.05em;margin-top:1rem;direction:rtl;">
            {answered} من ٥٦٧ عبارة تمت الإجابة عنها  ·  الصفحة {cp+1} من {total_pages}
        </div>
        <div class="progress-wrap">
            <div class="progress-fill" style="width:{pct}%"></div>
        </div>""", unsafe_allow_html=True)

        # Navigation
        c_prev, c_mid, c_next = st.columns([1,2,1])
        with c_prev:
            if cp > 0:
                if st.button("→ السابق"):
                    st.session_state.current_page -= 1; st.rerun()
        with c_next:
            if cp < total_pages - 1:
                if st.button("التالي ←"):
                    st.session_state.current_page += 1; st.rerun()

        # Submit on last page
        if cp == total_pages - 1:
            all_answered = len(st.session_state.responses) == 567
            client_name = st.session_state.get("client_name_saved","") or ""
            has_arabic  = any('\u0600' <= c <= '\u06ff' for c in client_name)
            gender_ar   = st.session_state.get("gender_saved","— اختر —")
            gender_en   = "Male" if gender_ar == "ذكر" else "Female" if gender_ar == "أنثى" else "Male"

            if not all_answered:
                st.markdown(f"""<div style="background:#FFF8F0;border-right:3px solid #E07B39;
                    border-left:none;padding:1rem 1.2rem;border-radius:4px 0 0 4px;
                    font-size:.88rem;color:#7A3D1A;margin:1rem 0;direction:rtl;text-align:right;">
                    ⚠ يرجى الإجابة على جميع العبارات قبل التسليم.
                    ({567 - answered} عبارة متبقية)
                </div>""", unsafe_allow_html=True)

            st.markdown('<div style="text-align:center;padding:2rem 0;">', unsafe_allow_html=True)
            submit = st.button("تسليم الاستبيان", disabled=not all_answered)
            st.markdown('</div>', unsafe_allow_html=True)

            if submit and has_arabic:
                st.markdown("""<div style="background:#FFF0F0;border-right:3px solid #D9534F;
                    border-left:none;padding:1rem 1.2rem;border-radius:4px 0 0 4px;
                    font-size:.92rem;color:#7A1A1A;margin:.5rem 0;direction:rtl;
                    text-align:right;font-weight:500;">
                    ⚠ يرجى كتابة اسمك باللغة الإنجليزية فقط. الأسماء المكتوبة بالعربية غير مقبولة.
                </div>""", unsafe_allow_html=True)

            if submit and all_answered and not has_arabic:
                with st.spinner("جاري معالجة إجاباتك..."):
                    scores  = compute_all_scores(st.session_state.responses, gender_en)
                    validity = check_validity(scores)
                    dob_v         = st.session_state.get("dob_saved","") or "Not provided"
                    nationality_v = st.session_state.get("nationality_saved","") or "Not provided"
                    referral_v    = st.session_state.get("referral_saved","") or "Not provided"
                    report  = generate_report(
                        client_name or "Anonymous",
                        st.session_state.get("age_saved","") or "N/A",
                        gender_en, scores, validity,
                        dob=dob_v, nationality=nationality_v, referral=referral_v
                    )
                    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    sn = (client_name or "anonymous").replace(" ","_").lower()
                    pdf_fname  = f"MMPI2_{sn}_{ts}.pdf"
                    docx_fname = f"MMPI2_AR_{sn}_{ts}.docx"
                    os.makedirs("reports", exist_ok=True)
                    pdf_path  = os.path.join("reports", pdf_fname)
                    docx_path = os.path.join("reports", docx_fname)

                    try:
                        create_pdf(pdf_path, client_name or "Anonymous",
                                   st.session_state.get("age_saved","") or "N/A",
                                   gender_en, scores, validity, report,
                                   dob=dob_v, nationality=nationality_v, referral=referral_v)
                    except Exception as e:
                        st.error(f"خطأ في إنشاء PDF: {e}"); st.stop()

                    try:
                        create_arabic_word_doc(
                            docx_path, client_name or "Anonymous",
                            st.session_state.get("age_saved","") or "N/A",
                            gender_ar or "—",
                            st.session_state.responses
                        )
                    except Exception as e:
                        st.warning(f"تحذير: فشل إنشاء ملف Word: {e}")
                        docx_path = None

                    email_error = None
                    try:
                        send_email(pdf_path, docx_path or pdf_path,
                                   client_name or "Anonymous",
                                   scores, pdf_fname,
                                   docx_fname if docx_path else pdf_fname)
                    except Exception as e:
                        email_error = str(e)

                    st.session_state.submitted   = True
                    st.session_state.email_error = email_error
                    st.rerun()
